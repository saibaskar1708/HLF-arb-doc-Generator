"""
generate_proceedings.py
───────────────────────
Generates Hearing Proceedings (.docx + .pdf) for each case in the
DataFormatted sheet of 'Lot 3 HLF Cases.xlsx'.

Each proceedings page contains:
  - Arbitrator header (BEFORE [ARB NAME], SOLE ARBITRATOR)
  - Case caption (case no, matter title, parties in bordered box)
  - "PROCEEDINGS FOR Nth MEETING HELD ON [date]"
  - Appearances & notice-service status
  - Adjournment details with next hearing date
  - QR code for the Google Meet link
  - Arbitrator name & signature (if image found)

Outputs
  Proceedings_Output/          – individual .docx per case
  HLF_Proceedings_Combined.docx – all cases merged into one DOCX
  HLF_Proceedings_Combined.pdf  – PDF version of the combined DOCX

Signature images
  Place PNG/JPG files in a  signatures/  sub-folder, named by arbitrator
  code (e.g.  signatures/CJS.png,  signatures/GB.jpg).
  Arbitrator codes are read from the Settings sheet.
  The "Arbitrator Signature" column (col 72) in the worksheet is also
  checked; if it contains a valid file path that file is used instead.

Required packages:
    pip install python-docx qrcode pillow openpyxl docx2pdf
"""

import os
import io
import sys
import copy
from datetime import datetime, date

# Ensure UTF-8 output on Windows consoles
if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

import openpyxl
import qrcode
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE_DIR       = os.path.dirname(os.path.abspath(__file__))
EXCEL_PATH     = os.path.join(BASE_DIR, "Lot 3 HLF Cases.xlsx")
SHEET_NAME     = "DataFormatted"
OUTPUT_DIR     = os.path.join(BASE_DIR, "Proceedings_Output")
SIGNATURES_DIR = os.path.join(BASE_DIR, "signatures")
COMBINED_DOCX  = os.path.join(BASE_DIR, "HLF_Proceedings_Combined.docx")
COMBINED_PDF   = os.path.join(BASE_DIR, "HLF_Proceedings_Combined.pdf")

# ── Column indices (0-based, matching iter_rows(values_only=True)) ─────────────
C_CONTRACT_NO          = 0
C_CONTRACT_DATE        = 1
C_BORROWER_NAME        = 16
C_BORROWER_ADDR        = 17
C_COBORROWER_NAME      = 18
C_COBORROWER_ADDR      = 19
C_GUARANTOR_NAME       = 20
C_GUARANTOR_ADDR       = 21
C_CASE_NO              = 27
C_ARB_NAME             = 29
C_ARB_ADDRESS          = 30
C_NOH_DATE             = 32
C_FIRST_HEARING_DATE   = 37
C_SECOND_HEARING_DATE  = 38
C_NOH_R1               = 45
C_NOH_R2               = 46
C_NOH_R3               = 47
C_MEETING_LINK         = 48
C_MEETING_TIMINGS      = 49
C_ARB_EMAIL            = 50
C_PLURAL               = 52   # "Respondent" / "Respondents"
C_ARB_SIGNATURE        = 72   # "Arbitrator Signature" (file path or #VALUE!)

# ── Page layout ────────────────────────────────────────────────────────────────
PAGE_W        = Cm(21)
PAGE_H        = Cm(29.7)
MARGIN_TOP    = Cm(1.5)
MARGIN_BOTTOM = Cm(1.5)
MARGIN_LEFT   = Cm(2.5)
MARGIN_RIGHT  = Cm(2.5)
CONTENT_W_CM  = 16.0   # 21 - 2*2.5

FONT_NAME  = 'Calibri'
FONT_SIZE  = 11        # pt

# ── Venue (fixed) ──────────────────────────────────────────────────────────────
VENUE = (
    "Sai and Sai Arbitration Centre, No.2, Diwan Bahadur Shanmugam Street, "
    "Kilpauk, Chennai- 600010. Ph: 91-44-48557697."
)

# ── Claimant (fixed) ───────────────────────────────────────────────────────────
CLAIMANT_TEXT = (
    "M/S. Hinduja Leyland Finance Ltd.,\n"
    "rep.by its Ms. Sumana B, Corporate legal,\n"
    "having its Corporate Office\n"
    "No.27A, Developed Industrial Estate, Guindy, Chennai-600 033."
)


# ── Utilities ──────────────────────────────────────────────────────────────────

def v(x):
    """Return stripped string; empty string for None."""
    return '' if x is None else str(x).strip()


def ordinal(n):
    n = int(n)
    suffix = 'th' if 11 <= (n % 100) <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th')
    return f"{n}{suffix}"


def fmt_date(d):
    """Format date/datetime/string as DD.MM.YYYY; return '' if empty."""
    if d is None or d == '':
        return ''
    if isinstance(d, (datetime, date)):
        return d.strftime('%d.%m.%Y')
    s = str(d).strip()
    for fmt in ('%Y-%m-%d %H:%M:%S', '%d-%m-%Y', '%d/%m/%Y', '%Y-%m-%d', '%d.%m.%Y'):
        try:
            return datetime.strptime(s, fmt).strftime('%d.%m.%Y')
        except ValueError:
            pass
    return s


def ordinal_word(n):
    """1 → 'First', 2 → 'Second', 3 → 'Third', 4 → 'Fourth'"""
    return {1: 'First', 2: 'Second', 3: 'Third', 4: 'Fourth'}.get(n, f"{ordinal(n)}")


def ordinal_upper(n):
    """1 → '1ST', 2 → '2ND', etc."""
    return ordinal(n).upper()


# ── python-docx helpers ────────────────────────────────────────────────────────

def set_para_fmt(para, align=None, space_before=0, space_after=3):
    pf = para.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)


def add_run(para, text, bold=False, italic=False, underline=False, size=FONT_SIZE):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    return run


def remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_outer_border(table, sz=4):
    """Add a thin outer border to a table with no inner borders (sz=4 → 0.5 pt)."""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    for edge in ('insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_cell_widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _set_vAlign(cell, val):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val)
    tcPr.append(vAlign)


def add_horizontal_rule(doc):
    """Add a full-width horizontal line paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_fmt(p, space_before=2, space_after=2)
    return p


# ── Signature / QR helpers ────────────────────────────────────────────────────

def make_sig_transparent(img_path, threshold=240):
    """Remove white background from signature image; return high-res BytesIO PNG."""
    img = Image.open(img_path).convert('RGBA')
    data = img.getdata()
    new_pixels = [
        (r, g, b, 0) if r >= threshold and g >= threshold and b >= threshold else (r, g, b, a)
        for r, g, b, a in data
    ]
    img.putdata(new_pixels)
    buf = io.BytesIO()
    img.save(buf, format='PNG', dpi=(300, 300))
    buf.seek(0)
    return buf


def generate_qr(url):
    """Generate a high-res QR code PNG from url; return BytesIO."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=20,
        border=4,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color='black', back_color='white')
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


# ── Arbitrator signature lookup ───────────────────────────────────────────────

def load_arb_code_map(excel_path):
    """Return dict: arbitrator_name_lower → arbitrator_code (from Settings sheet)."""
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        ws = wb['Settings']
        arb_map = {}
        for row in range(4, ws.max_row + 1):
            code = ws.cell(row, 1).value
            name = ws.cell(row, 2).value
            if code and name:
                arb_map[str(name).lower().strip()] = str(code).strip()
        wb.close()
        return arb_map
    except Exception:
        return {}


def find_sig_path(arb_name, arb_code_map, sig_dir):
    """
    Find a signature image file.
    1. Look in sig_dir for {code}.png/.jpg/.jpeg/.bmp
    2. Fall back to {arb_name_cleaned}.png etc.
    Returns file path or None.
    """
    code = arb_code_map.get(arb_name.lower().strip(), '')
    candidates = []
    if code:
        for ext in ('.png', '.jpg', '.jpeg', '.bmp'):
            candidates.append(os.path.join(sig_dir, f"{code}{ext}"))
            candidates.append(os.path.join(sig_dir, f"{code.lower()}{ext}"))
    safe_name = ''.join(c for c in arb_name if c.isalnum() or c in ' _')
    safe_name = safe_name.replace(' ', '_').strip('_')
    for ext in ('.png', '.jpg', '.jpeg', '.bmp'):
        candidates.append(os.path.join(sig_dir, f"{safe_name}{ext}"))

    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


# ── Notice service sentence builder ──────────────────────────────────────────

def build_notice_sentence(noh_date_str, r1_status, r2_status, r3_status, respondents):
    """
    Build the notice service sentence.
    respondents: list of (num, name, address) tuples (length 1, 2 or 3)
    r1/r2/r3_status: string like 'Served', 'Returned', 'Not Served', or None → 'Served'
    """
    def status(s):
        return v(s) if v(s) else 'Served'

    n = len(respondents)
    if n == 0:
        return ''

    parts = []
    parts.append(f"The Notice dated {noh_date_str} sent to the {ordinal_word(1)} Respondent is {status(r1_status)}")
    if n >= 2:
        parts.append(f"the {ordinal_word(2)} Respondent is {status(r2_status)}")
    if n >= 3:
        parts.append(f"the {ordinal_word(3)} Respondent is {status(r3_status)}")

    if len(parts) == 1:
        return parts[0] + '.'
    elif len(parts) == 2:
        return parts[0] + ' and ' + parts[1] + '.'
    else:
        return parts[0] + ', ' + parts[1] + ' and ' + parts[2] + '.'


# ── Meeting number helper ─────────────────────────────────────────────────────

def determine_meeting_number(first_hearing_date, second_hearing_date):
    """
    If only First Hearing Date → 1st meeting proceedings.
    If Second Hearing Date is also present → 2nd meeting proceedings.
    """
    if v(second_hearing_date):
        return 2, fmt_date(second_hearing_date)
    return 1, fmt_date(first_hearing_date)


# ── Core document builder ─────────────────────────────────────────────────────

def build_proceedings(row_data, sig_image_path=None, sig_col_idx=None):
    """Build and return a proceedings Document for one data row."""

    # ── Extract values ────────────────────────────────────────────────────────
    contract_no   = v(row_data[C_CONTRACT_NO])
    contract_date = fmt_date(row_data[C_CONTRACT_DATE])
    borrower_name = v(row_data[C_BORROWER_NAME])
    borrower_addr = v(row_data[C_BORROWER_ADDR])
    coborrower_name = v(row_data[C_COBORROWER_NAME])
    coborrower_addr = v(row_data[C_COBORROWER_ADDR])
    guarantor_name  = v(row_data[C_GUARANTOR_NAME])
    guarantor_addr  = v(row_data[C_GUARANTOR_ADDR])

    # Resolve "same as above" address shortcuts
    _SAME = {'same as above', 'same as borrower', 'same', '-', ''}
    if coborrower_addr.lower() in _SAME:
        coborrower_addr = borrower_addr
    if guarantor_addr.lower() in _SAME:
        guarantor_addr = borrower_addr

    case_no         = v(row_data[C_CASE_NO])
    arb_name        = v(row_data[C_ARB_NAME])
    noh_date_str    = fmt_date(row_data[C_NOH_DATE])
    first_hearing   = row_data[C_FIRST_HEARING_DATE]
    second_hearing  = row_data[C_SECOND_HEARING_DATE]
    noh_r1          = v(row_data[C_NOH_R1])
    noh_r2          = v(row_data[C_NOH_R2])
    noh_r3          = v(row_data[C_NOH_R3])
    meeting_link    = v(row_data[C_MEETING_LINK])
    meeting_timings = v(row_data[C_MEETING_TIMINGS])
    arb_email       = v(row_data[C_ARB_EMAIL])
    plural          = v(row_data[C_PLURAL]) if len(row_data) > C_PLURAL else 'Respondents'

    # Check for inline signature path from "Arbitrator Signature" column
    # Skip Excel error values (e.g. #VALUE!)
    if sig_col_idx is not None and sig_image_path is None:
        sig_val = v(row_data[sig_col_idx]) if len(row_data) > sig_col_idx else ''
        if sig_val and not sig_val.startswith('#') and os.path.isfile(sig_val):
            sig_image_path = sig_val

    # Meeting details
    meeting_num, held_on_date = determine_meeting_number(first_hearing, second_hearing)
    # Next meeting date (for adjournment text)
    next_date = fmt_date(second_hearing) if (meeting_num == 1 and v(second_hearing)) else '_____'

    # Build respondents list: [(number_str, name, address), ...]
    respondents = [('1.', borrower_name, borrower_addr)]
    if coborrower_name:
        respondents.append(('2.', coborrower_name, coborrower_addr))
    if guarantor_name:
        num = '3.' if coborrower_name else '2.'
        respondents.append((num, guarantor_name, guarantor_addr))

    resp_label = plural if plural else ('Respondents' if len(respondents) > 1 else 'Respondent')
    resp_count = len(respondents)

    if resp_count == 1:
        matter_parties = borrower_name
    elif resp_count == 2:
        matter_parties = f"{borrower_name} and another"
    else:
        matter_parties = f"{borrower_name} and others"

    matter_title = (
        f"In the matter of Arbitration & Conciliation Act 1996 "
        f"And  In the matter of disputes between M/S. Hinduja Leyland Finance Ltd., "
        f"and {matter_parties}  arising under the Loan Agreement   "
        f"No.{contract_no}  Dated {contract_date}"
    )

    if meeting_timings:
        timings_fmt = meeting_timings.replace(' - ', ' \u2013 ')
    else:
        timings_fmt = '_____ \u2013 _____'

    adj_text = (
        f"The matter is adjourned to {next_date} from {timings_fmt} "
        f"at the venue {VENUE} Email: {arb_email}. "
        f"Alternatively, you may choose to attend the proceedings via video conferencing. "
        f"The video conference details are stated below"
    )

    notice_sentence = build_notice_sentence(
        noh_date_str, noh_r1, noh_r2, noh_r3, respondents
    )

    appearances_text = (
        f"The authorised representative of the Claimant appeared, and none appeared for "
        f"the {resp_label}. {notice_sentence}"
    )

    # ── Create document ───────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = PAGE_W
    sec.page_height   = PAGE_H
    sec.top_margin    = MARGIN_TOP
    sec.bottom_margin = MARGIN_BOTTOM
    sec.left_margin   = MARGIN_LEFT
    sec.right_margin  = MARGIN_RIGHT

    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE)

    # ── 0. BEFORE [Arbitrator Name] header ────────────────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    add_run(p, 'BEFORE', bold=True, size=10)

    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    add_run(p, arb_name.upper(), bold=True, size=13)

    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
    add_run(p, 'SOLE ARBITRATOR', bold=True, size=10)

    # ── Horizontal rule ───────────────────────────────────────────────────────
    add_horizontal_rule(doc)

    # ── 1. Case number ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=5, space_after=3)
    add_run(p, case_no, bold=True, size=12)

    # ── 2. Matter title ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=5)
    add_run(p, matter_title)

    # ── 3. Parties borderless table (Claimant / AND / Respondents) ───────────
    LEFT_W  = CONTENT_W_CM - 3.5   # 12.5 cm
    RIGHT_W = 3.5                   # 3.5 cm

    pt = doc.add_table(rows=3, cols=2)
    remove_table_borders(pt)

    # Row 0: Claimant
    claimant_left  = pt.rows[0].cells[0]
    claimant_right = pt.rows[0].cells[1]
    claimant_left.width  = Cm(LEFT_W)
    claimant_right.width = Cm(RIGHT_W)
    _set_vAlign(claimant_right, 'bottom')

    lp = claimant_left.paragraphs[0]
    set_para_fmt(lp, space_before=2, space_after=2)
    add_run(lp, CLAIMANT_TEXT)

    rp = claimant_right.paragraphs[0]
    set_para_fmt(rp, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=2)
    add_run(rp, '\u2026Claimant')

    # Row 1: AND (merged across both columns)
    and_cell = pt.rows[1].cells[0].merge(pt.rows[1].cells[1])
    and_p = and_cell.paragraphs[0]
    set_para_fmt(and_p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)
    add_run(and_p, 'AND', bold=True)

    # Row 2: Respondents — address compressed to one line to save vertical space
    resp_left  = pt.rows[2].cells[0]
    resp_right = pt.rows[2].cells[1]
    resp_left.width  = Cm(LEFT_W)
    resp_right.width = Cm(RIGHT_W)
    _set_vAlign(resp_right, 'bottom')

    resp_p = resp_left.paragraphs[0]
    set_para_fmt(resp_p, space_before=2, space_after=2)
    for i, (num, name, addr) in enumerate(respondents):
        prefix = '\n' if i > 0 else ''
        addr_parts = [l.strip() for l in addr.strip().replace('\r', '').split('\n') if l.strip()]
        addr_text  = ', '.join(addr_parts)   # single line, comma-separated
        full_block = f"{prefix}{num} {name}"
        if addr_text:
            full_block += f", {addr_text}"
        add_run(resp_p, full_block)

    rl = resp_right.paragraphs[0]
    set_para_fmt(rl, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=2)
    add_run(rl, f'\u2026{resp_label}')

    # ── 4. Proceedings heading ─────────────────────────────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    add_run(p, f'PROCEEDINGS FOR {ordinal_upper(meeting_num)} MEETING HELD ON {held_on_date}',
            bold=True, size=12)

    # ── 5. Appearances + notices ──────────────────────────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=4)
    add_run(p, appearances_text)

    # ── 6. Adjournment ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6)
    add_run(p, adj_text)

    # ── 7. Google Meet section ────────────────────────────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=1)
    add_run(p, 'Google Meet joining info', bold=True)

    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=1)
    add_run(p, f'Video call link: {meeting_link}')

    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=2)
    add_run(p, 'Or scan the QR Code below to join the meeting')

    # QR code image — high-res (box_size=20), displayed at 2.5 cm for sharpness
    if meeting_link:
        qr_buf  = generate_qr(meeting_link)
        qr_para = doc.add_paragraph()
        set_para_fmt(qr_para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=8)
        qr_para.add_run().add_picture(qr_buf, width=Cm(2.5))

    # ── 8. Arbitrator signature block ─────────────────────────────────────────
    if sig_image_path:
        sig_para = doc.add_paragraph()
        set_para_fmt(sig_para, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
        sig_para.add_run().add_picture(make_sig_transparent(sig_image_path), width=Cm(2.5))

    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
    add_run(p, arb_name, bold=True)

    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
    add_run(p, 'Arbitrator')

    return doc


# ── Document merger ────────────────────────────────────────────────────────────

def merge_docs(doc_bufs):
    """Merge a list of BytesIO DOCX buffers into one Document with page breaks."""
    combined = Document(doc_bufs[0])
    for buf in doc_bufs[1:]:
        src  = Document(buf)
        body = combined.element.body
        # Hard page break before appending next document
        br_p  = OxmlElement('w:p')
        br_r  = OxmlElement('w:r')
        br_el = OxmlElement('w:br')
        br_el.set(qn('w:type'), 'page')
        br_r.append(br_el)
        br_p.append(br_r)
        last = list(body)[-1]
        if last.tag == qn('w:sectPr'):
            last.addprevious(br_p)
        else:
            body.append(br_p)
        # Copy body elements (skip trailing sectPr)
        for el in list(src.element.body):
            if el.tag == qn('w:sectPr'):
                continue
            el_copy = copy.deepcopy(el)
            last = list(body)[-1]
            if last.tag == qn('w:sectPr'):
                last.addprevious(el_copy)
            else:
                body.append(el_copy)
    return combined


# ── Main batch processor ──────────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Load arbitrator code map for signature lookup
    arb_code_map = load_arb_code_map(EXCEL_PATH)

    # Load worksheet data
    wb = openpyxl.load_workbook(EXCEL_PATH, read_only=True, data_only=True)
    ws = wb[SHEET_NAME]
    all_rows = list(ws.iter_rows(values_only=True))
    wb.close()

    headers  = list(all_rows[0]) if all_rows else []
    data     = all_rows[1:]

    # Check if an "Arbitrator Signature" column exists (col 72 in Lot 3)
    sig_col_idx = None
    for i, h in enumerate(headers):
        if h and 'signature' in str(h).lower():
            sig_col_idx = i
            print(f"  Found signature column: '{h}' (index {i})")
            break

    total     = 0
    generated = 0
    skipped   = []
    doc_bufs  = []

    print()
    print("=" * 65)
    print("  HLF Hearing Proceedings Generator")
    print("=" * 65)
    print(f"  Excel   : {EXCEL_PATH}")
    print(f"  Sheet   : {SHEET_NAME}")
    print(f"  Output  : {OUTPUT_DIR}")
    print()

    for row_data in data:
        if all(c is None for c in row_data):
            continue

        contract_no   = v(row_data[C_CONTRACT_NO])
        borrower_name = v(row_data[C_BORROWER_NAME])
        case_no       = v(row_data[C_CASE_NO])
        first_hearing = row_data[C_FIRST_HEARING_DATE]
        meeting_link  = v(row_data[C_MEETING_LINK])
        arb_name      = v(row_data[C_ARB_NAME])

        total += 1

        # Skip rows with no hearing date (not scheduled yet)
        if not first_hearing:
            skipped.append({
                'contract_no':   contract_no,
                'borrower_name': borrower_name,
                'reason':        'Missing: First Hearing Date',
            })
            continue

        # Skip rows with no meeting link (needed for QR code)
        if not meeting_link:
            skipped.append({
                'contract_no':   contract_no,
                'borrower_name': borrower_name,
                'reason':        'Missing: Meeting Link',
            })
            continue

        try:
            # Resolve signature image from signatures/ folder
            sig_path = None
            if os.path.isdir(SIGNATURES_DIR):
                sig_path = find_sig_path(arb_name, arb_code_map, SIGNATURES_DIR)

            doc = build_proceedings(row_data, sig_image_path=sig_path, sig_col_idx=sig_col_idx)

            # Save individual file
            safe_case     = case_no.replace('/', '_').replace('\\', '_')
            safe_contract = contract_no.replace('/', '_').replace('\\', '_')
            fname    = f"Proceedings_{safe_case}_{safe_contract}.docx"
            filepath = os.path.join(OUTPUT_DIR, fname)
            doc.save(filepath)

            # Keep in-memory buffer for combined file
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            doc_bufs.append(buf)

            generated += 1
            print(f"  [OK] {fname}")

        except Exception as exc:
            import traceback
            skipped.append({
                'contract_no':   contract_no,
                'borrower_name': borrower_name,
                'reason':        f"Error: {exc}",
            })
            traceback.print_exc()

    # ── Combined DOCX ─────────────────────────────────────────────────────────
    if doc_bufs:
        print()
        print(f"  Merging {len(doc_bufs)} document(s) into combined DOCX...")
        combined_doc = merge_docs(doc_bufs)
        combined_doc.save(COMBINED_DOCX)
        print(f"  [OK] Combined DOCX: {COMBINED_DOCX}")

        # ── Convert to PDF ────────────────────────────────────────────────────
        try:
            from docx2pdf import convert
            print(f"  Converting to PDF...")
            convert(COMBINED_DOCX, COMBINED_PDF)
            print(f"  [OK] Combined PDF : {COMBINED_PDF}")
        except ImportError:
            print("  docx2pdf not installed - skipping PDF conversion.")
            print("     Run:  pip install docx2pdf")
        except Exception as pdf_exc:
            print(f"  PDF conversion failed: {pdf_exc}")
            print("     Ensure Microsoft Word is installed and accessible.")
    else:
        print()
        print("  No documents generated - nothing to combine.")

    # ── Summary ───────────────────────────────────────────────────────────────
    print()
    print("=" * 65)
    print("  BATCH SUMMARY - Hearing Proceedings Generator")
    print("=" * 65)
    print(f"  Total rows processed  : {total}")
    print(f"  Documents generated   : {generated}")
    print(f"  Rows skipped          : {len(skipped)}")
    if skipped:
        print()
        print(f"  {'Contract No':<25} {'Borrower':<28} Reason")
        print(f"  {'-'*25} {'-'*28} {'-'*28}")
        for s in skipped:
            print(f"  {s['contract_no']:<25} {s['borrower_name']:<28} {s['reason']}")
    print("=" * 65)
    print()


if __name__ == '__main__':
    main()
