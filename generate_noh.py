"""
generate_noh.py
───────────────
Generates Notice of Hearing (NOH) Word documents (.docx) for each case
in the DataFormatted sheet of 'Lot 4 HLF Cases.xlsx'.

Each document is 2 pages:
  Page 1 – Notice of Hearing  (fits one A4 page)
  Page 2 – Disclosure U/S.12(1) with Sixth Schedule table

Output files are saved to the NOH_Output/ folder.
Rows missing mandatory fields (NOH Date, Meeting Link) are skipped and
reported in the end-of-run summary.

Required packages:
    pip install python-docx qrcode pillow openpyxl
"""

import os
import io
import copy
from datetime import datetime, date

import openpyxl
import qrcode
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
EXCEL_PATH = os.path.join(BASE_DIR, "Lot 4 HLF Cases.xlsx")
SHEET_NAME = "DataFormatted"
OUTPUT_DIR = os.path.join(BASE_DIR, "NOH_Output")

# ── Column indices (0-based) ───────────────────────────────────────────────────
C_CONTRACT_NO        = 0
C_CONTRACT_DATE      = 1
C_CONTRACT_STATUS    = 2
C_BORROWER_NAME      = 16
C_BORROWER_ADDR      = 17
C_COBORROWER_NAME    = 18
C_COBORROWER_ADDR    = 19
C_GUARANTOR_NAME     = 20
C_GUARANTOR_ADDR     = 21
C_REF_LETTER_DATE    = 26
C_CASE_NO            = 27
C_ARB_APPT_DATE      = 28
C_ARB_NAME           = 29
C_ARB_ADDRESS        = 30
C_ARB_ACCEPT_DATE    = 31
C_NOH_DATE           = 32
C_FIRST_HEARING_DATE = 37
C_MEETING_LINK       = 48
C_MEETING_TIMINGS    = 49
C_ARB_EMAIL          = 50
C_REFER_BORROWER     = 53
C_ARB_EXPERIENCE     = 71

# ── Page layout constants ──────────────────────────────────────────────────────
PAGE_W        = Cm(21)
PAGE_H        = Cm(29.7)
MARGIN_TOP    = Cm(1.2)
MARGIN_BOTTOM = Cm(1.2)
MARGIN_LEFT   = Cm(2.0)
MARGIN_RIGHT  = Cm(2.0)
CONTENT_W_CM  = 17.0   # 21 - 2*2

FONT_NAME     = 'Times New Roman'
FONT_SIZE     = 10     # pt

# ── Utility helpers ────────────────────────────────────────────────────────────

def v(x):
    """Return stripped string; empty string for None."""
    return '' if x is None else str(x).strip()


def ordinal(n):
    """1 → '1st', 2 → '2nd', 3 → '3rd', 4 → '4th', …"""
    n = int(n)
    suffix = 'th' if 11 <= (n % 100) <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th')
    return f"{n}{suffix}"


def fmt_date(d):
    """Format date/datetime/string as DD.MM.YYYY; return '' if empty."""
    if d is None or d == '':
        return ''
    if isinstance(d, (datetime, date)):
        return d.strftime('%d.%m.%Y')
    s = str(d).strip()
    for fmt in ('%Y-%m-%d %H:%M:%S', '%d-%m-%Y', '%d/%m/%Y', '%Y-%m-%d', '%d.%m.%Y'):
        try:
            return datetime.strptime(s, fmt).strftime('%d.%m.%Y')
        except ValueError:
            pass
    return s


def fmt_date_long(d):
    """Format date as 'Dated the 9th day of February 2026'."""
    if d is None or d == '':
        return ''
    if isinstance(d, (datetime, date)):
        return f"Dated the {ordinal(d.day)} day of {d.strftime('%B %Y')}"
    s = str(d).strip()
    for fmt in ('%Y-%m-%d %H:%M:%S', '%d-%m-%Y', '%d/%m/%Y', '%Y-%m-%d', '%d.%m.%Y'):
        try:
            dt = datetime.strptime(s, fmt)
            return f"Dated the {ordinal(dt.day)} day of {dt.strftime('%B %Y')}"
        except ValueError:
            pass
    return s


def party_ordinals_text(resp_count):
    """Return ordinal range text based on number of respondents."""
    if resp_count == 1:
        return "the 1st and 2nd"
    elif resp_count == 2:
        return "the 1st, 2nd and 3rd"
    else:
        return "the 1st, 2nd, 3rd and 4th"


def resp_label(resp_count):
    return "Respondents" if resp_count > 1 else "Respondent"


def clip_address(addr, max_lines=3, chars_per_line=76):
    """Reflow address text so it fits within max_lines lines of chars_per_line chars."""
    if not addr:
        return addr
    # Tokenise on existing newlines and commas
    raw = addr.replace('\r', '').replace('\n', ',')
    tokens = [t.strip() for t in raw.split(',') if t.strip()]
    lines, current = [], ''
    for token in tokens:
        sep = ', ' if current else ''
        candidate = current + sep + token
        if len(candidate) <= chars_per_line:
            current = candidate
        else:
            if current:
                lines.append(current)
                if len(lines) >= max_lines:
                    current = ''
                    break
            current = token
    if current and len(lines) < max_lines:
        lines.append(current)
    return '\n'.join(lines)


# ── python-docx helpers ────────────────────────────────────────────────────────

def set_para_fmt(para, align=None, space_before=0, space_after=2):
    pf = para.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)


def add_run(para, text, bold=False, underline=False, size=FONT_SIZE):
    run = para.add_run(text)
    run.bold      = bold
    run.underline = underline
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    return run


def add_bottom_border(paragraph):
    """Draw a horizontal rule below a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def remove_table_borders(table):
    """Remove all visible borders from a table."""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_cell_widths(table, widths_cm):
    """Set each column width per cell (more reliable than table.columns)."""
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _set_vAlign(cell, val):
    """Set vertical alignment of a table cell (top / center / bottom)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val)
    tcPr.append(vAlign)


def make_sig_transparent(img_path, threshold=240):
    """Open a signature image, remove white/near-white background, return BytesIO PNG."""
    img = Image.open(img_path).convert('RGBA')
    pixels = list(img.getdata())
    new_pixels = [
        (r, g, b, 0) if r >= threshold and g >= threshold and b >= threshold else (r, g, b, a)
        for r, g, b, a in pixels
    ]
    img.putdata(new_pixels)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def generate_qr(url, size_cm=3.0):
    """Generate a QR code PNG from url; return BytesIO."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=10,
        border=4,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color='black', back_color='white')
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


# ── Document builder ───────────────────────────────────────────────────────────

def build_noh(row_data, sig_image_path=None):
    """Build and return a 2-page NOH Document for one data row."""

    # ── Extract & format values ──────────────────────────────────────────────
    contract_no     = v(row_data[C_CONTRACT_NO])
    contract_date   = fmt_date(row_data[C_CONTRACT_DATE])
    contract_status = v(row_data[C_CONTRACT_STATUS]).upper()
    borrower_name   = v(row_data[C_BORROWER_NAME])
    borrower_addr   = v(row_data[C_BORROWER_ADDR])
    coborrower_name = v(row_data[C_COBORROWER_NAME])
    coborrower_addr = v(row_data[C_COBORROWER_ADDR])
    guarantor_name  = v(row_data[C_GUARANTOR_NAME])
    guarantor_addr  = v(row_data[C_GUARANTOR_ADDR])

    # Resolve "same as above" / "same as borrower" placeholders to the actual address
    _SAME_AS = {'same as above', 'same as borrower', 'same', '-', ''}
    if coborrower_addr.lower() in _SAME_AS:
        coborrower_addr = borrower_addr
    if guarantor_addr.lower() in _SAME_AS:
        guarantor_addr = borrower_addr
    honorific       = v(row_data[C_REFER_BORROWER])
    arb_name        = v(row_data[C_ARB_NAME])
    arb_address     = v(row_data[C_ARB_ADDRESS])
    arb_email       = v(row_data[C_ARB_EMAIL])
    arb_experience  = v(row_data[C_ARB_EXPERIENCE])
    case_no         = v(row_data[C_CASE_NO])
    arb_appt_date   = fmt_date(row_data[C_ARB_APPT_DATE])
    ref_letter_date = fmt_date(row_data[C_REF_LETTER_DATE])
    arb_accept_date = fmt_date(row_data[C_ARB_ACCEPT_DATE])
    dated_line      = fmt_date_long(row_data[C_NOH_DATE])
    first_hearing   = fmt_date(row_data[C_FIRST_HEARING_DATE])
    meeting_link    = v(row_data[C_MEETING_LINK])
    meeting_timings = v(row_data[C_MEETING_TIMINGS])

    # Respondents list (always starts from "2.")
    # Clip each address to 3 lines to prevent page overflow
    respondents = [('2.', borrower_name, clip_address(borrower_addr))]
    if coborrower_name:
        respondents.append(('3.', coborrower_name, clip_address(coborrower_addr)))
    if guarantor_name:
        num = '4.' if coborrower_name else '3.'
        respondents.append((num, guarantor_name, clip_address(guarantor_addr)))
    resp_count = len(respondents)

    # Conditional last sentence (Contract Status L or R)
    if contract_status in ('L', 'R'):
        last_sentence = (
            "The Claimant has filed their Claim Statement along with the affidavit and the "
            "petition under the Arbitration and Conciliation Act 1996, which are appended herewith."
        )
    else:
        last_sentence = (
            "The Claimant has filed their Claim Statement, which is appended herewith."
        )

    # ── Create document ──────────────────────────────────────────────────────
    doc = Document()

    # Page setup: A4, narrow margins
    sec = doc.sections[0]
    sec.page_width    = PAGE_W
    sec.page_height   = PAGE_H
    sec.top_margin    = MARGIN_TOP
    sec.bottom_margin = MARGIN_BOTTOM
    sec.left_margin   = MARGIN_LEFT
    sec.right_margin  = MARGIN_RIGHT

    # Default paragraph style
    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE)

    # ════════════════════════════════════════════════════════════════════════
    # PAGE 1 — NOTICE OF HEARING
    # ════════════════════════════════════════════════════════════════════════

    # ── Header letterhead: borderless 2-col table ───────────────────────────
    hdr_tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(hdr_tbl)
    set_cell_widths(hdr_tbl, [11.0, 6.0])

    # Left cell: Arbitrator Name (bold) + line break + ARBITRATOR
    lp = hdr_tbl.cell(0, 0).paragraphs[0]
    set_para_fmt(lp, space_before=0, space_after=0)
    add_run(lp, arb_name, bold=True)
    lp.add_run('\n').font.size = Pt(FONT_SIZE)
    add_run(lp, 'ARBITRATOR')

    # Right cell: address, left-aligned (narrower column forces 2-3 line wrap)
    rp = hdr_tbl.cell(0, 1).paragraphs[0]
    set_para_fmt(rp, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0)
    add_run(rp, arb_address)

    # Horizontal divider line
    hr = doc.add_paragraph()
    set_para_fmt(hr, space_before=2, space_after=2)
    add_bottom_border(hr)

    # ── Section 1: Tribunal heading (centered, bold) ─────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=0)
    add_run(p, f'BEFORE THE ARBITRATOR & ADVOCATE {arb_name}', bold=True)

    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=3)
    add_run(p, f'CLAIM PETITION NO. {case_no}', bold=True)

    # ── Section 2: Parties ───────────────────────────────────────────────────
    # Claimant
    pt = doc.add_table(rows=1, cols=2)
    remove_table_borders(pt)
    set_cell_widths(pt, [13.5, 3.5])

    cp = pt.cell(0, 0).paragraphs[0]
    set_para_fmt(cp, space_before=0, space_after=0)
    add_run(cp,
        '1. M/S. Hinduja Leyland Finance Ltd.,\n'
        'Rep. By Its Authorized Representative Ms. Sumana B - Corporate Legal\n'
        'Having Its Corporate Office\n'
        'No.27a, Developed Industrial Estate,\n'
        'Guindy, Chennai-600 032.'
    )

    clp_cell = pt.cell(0, 1)
    _set_vAlign(clp_cell, 'bottom')
    clp = clp_cell.paragraphs[0]
    set_para_fmt(clp, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
    add_run(clp, '.... Claimant')

    # Vs  (tighten spacing when 3 respondents to avoid page overflow)
    tight = (resp_count == 3)
    vs = doc.add_paragraph()
    set_para_fmt(vs, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=1 if tight else 2, space_after=1 if tight else 2)
    add_run(vs, 'Vs')

    # Respondents
    rt = doc.add_table(rows=1, cols=2)
    remove_table_borders(rt)
    set_cell_widths(rt, [13.5, 3.5])

    resp_lines = []
    for i, (num, name, addr) in enumerate(respondents):
        resp_lines.append(f'{num} {name}\n{addr}')
    # Use single blank line between respondents when there are 3 (saves space)
    resp_text = '\n\n'.join(resp_lines) if resp_count < 3 else '\n'.join(resp_lines)

    rlp = rt.cell(0, 0).paragraphs[0]
    set_para_fmt(rlp, space_before=0, space_after=0)
    add_run(rlp, resp_text)

    rrp_cell = rt.cell(0, 1)
    _set_vAlign(rrp_cell, 'bottom')
    rrp = rrp_cell.paragraphs[0]
    set_para_fmt(rrp, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
    add_run(rrp, f'.... {resp_label(resp_count)}')

    # ── Section 3: Salutation, Ref, Sub ─────────────────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=2 if tight else 4, space_after=2)
    add_run(p, 'Sir/Madam,')

    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=2)
    add_run(p, 'Ref: ', bold=True)
    add_run(p,
        f'Letter dated {arb_appt_date} from Sai & Sai Arbitration Centre \u2013 '
        f'Appointment of Arbitrator \u2013 Arbitration in the matter of dispute(s) between '
        f'M/S. Hinduja Leyland Finance Ltd. vs. {honorific} {borrower_name} in respect of '
        f'Loan Account No. {contract_no} dated {contract_date}.'
    )

    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
    add_run(p, 'Sub: Notice of Hearing', bold=True)

    # ── Section 4: Body paragraph ────────────────────────────────────────────
    body_text = (
        f'Sai & Sai Arbitration Centre through its letter dated {arb_appt_date} had nominated '
        f'and appointed me as Arbitrator to arbitrate on the disputes/claim arisen between '
        f'{party_ordinals_text(resp_count)} of you which was in furtherance to Letter dated '
        f'{ref_letter_date} from the first of you. I hereby accept my appointment as Arbitrator '
        f'through the letter dated {arb_accept_date}. There are no circumstances exist that give '
        f'rise to justifiable doubts as to my independence or impartiality in resolving the dispute '
        f'referred. Declaration under Section 12(1) of the Arbitration and Conciliation Act, 1996 '
        f'\u2013 as per Sixth Schedule is also annexed herewith. {last_sentence}'
    )
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=2)
    add_run(p, body_text)

    # ── Section 5: Hearing notice ────────────────────────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=2)
    add_run(p,
        f'Take notice that the above matter stands posted for hearing on {first_hearing} '
        f'between {meeting_timings}. You have the option to appear either in person or through '
        f'your authorized representative at:'
    )

    # Section 5a: Venue address
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=2)
    add_run(p,
        f'Sai and Sai Arbitration Centre,\n'
        f'No.2, Diwan Bahadur Shanmugam Street,\n'
        f'Kilpauk, Chennai- 600010,\n'
        f'e-mail: {arb_email}  Phone: +91 44 48557697.'
    )

    # ── Section 6: Video conferencing block ─────────────────────────────────
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=2)
    add_run(p,
        'Alternatively, you may choose to attend the proceedings via video conferencing '
        'for which the details are stated below'
    )

    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=1)
    add_run(p, 'Google Meet Info', bold=True, underline=True)

    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=1)
    add_run(p, f'Video call link: {meeting_link}')

    # ── Section 7: 3-column table — QR | Date | Signature ───────────────────
    # Col 1: scan instruction + QR code (left-aligned)
    # Col 2: dated line (top-aligned, centered)
    # Col 3: Arbitrator signature block (right-aligned)
    bottom_tbl = doc.add_table(rows=1, cols=3)
    remove_table_borders(bottom_tbl)
    set_cell_widths(bottom_tbl, [5.0, 6.5, 5.5])

    # Col 1 — QR code
    qr_cell = bottom_tbl.cell(0, 0)
    qp = qr_cell.paragraphs[0]
    set_para_fmt(qp, space_before=0, space_after=2)
    add_run(qp, 'Or scan the QR Code below to join the meeting')
    qr_pic_para = qr_cell.add_paragraph()
    set_para_fmt(qr_pic_para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0)
    qr_pic_para.add_run().add_picture(generate_qr(meeting_link), width=Cm(3))

    # Col 2 — Dated line (vertically centered, horizontally centered)
    date_cell = bottom_tbl.cell(0, 1)
    _set_vAlign(date_cell, 'center')
    dp = date_cell.paragraphs[0]
    set_para_fmt(dp, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    add_run(dp, dated_line)

    # Col 3 — Signature block (bottom-aligned, right-aligned)
    sig_cell = bottom_tbl.cell(0, 2)
    _set_vAlign(sig_cell, 'bottom')

    if sig_image_path:
        sig_img_para = sig_cell.paragraphs[0]
        set_para_fmt(sig_img_para, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
        sig_img_para.add_run().add_picture(
            make_sig_transparent(sig_image_path), width=Cm(2.8)
        )
        sp = sig_cell.add_paragraph()
    else:
        sp = sig_cell.paragraphs[0]

    set_para_fmt(sp, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
    add_run(sp, f'{arb_name}\n(ARBITRATOR)')

    # ════════════════════════════════════════════════════════════════════════
    # PAGE 2 — DISCLOSURE U/S.12(1) WITH SIXTH SCHEDULE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
    add_run(p,
        'DISCLOSURE PROVIDED U/S.12 (1) READ WITH SIXTH SCHEDULE OF THE '
        'ARBITRATION AND CONCILIATION ACT, 1996',
        bold=True
    )

    # Disclosure table (3 cols, header + 6 data rows)
    disc_tbl = doc.add_table(rows=7, cols=3)
    disc_tbl.style = 'Table Grid'
    set_cell_widths(disc_tbl, [1.5, 9.0, 6.5])

    # Header row
    for cell, text in zip(disc_tbl.rows[0].cells, ['Sr. No.', 'Particulars', 'Details']):
        p = cell.paragraphs[0]
        set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)
        add_run(p, text, bold=True)

    # Data rows
    disc_data = [
        ('1', 'Name of the Arbitrator',
              arb_name),
        ('2', 'Contact Details',
              arb_address),
        ('3', 'Prior experience (including experience with Arbitrations)',
              arb_experience),
        ('4', 'Number of on-going arbitrations',
              ''),
        ('5', ('Circumstances disclosing any past or present relationship with or interest in '
               'any of the parties or in relation to the subject matter in dispute whether '
               'financial, business professional or any other kind which is likely to give rise '
               'to justifiable doubts as to the independence or impartiality'),
              'No vested interest with any of the parties'),
        ('6', ('Circumstances which are likely to affect the ability to devote sufficient time '
               'to the arbitration and in particular the ability to finish the entire arbitration '
               'within twelve months'),
              'No adverse circumstances to affect the ability to devote sufficient time in '
              'finishing the proceedings as stipulated.'),
    ]

    for i, (sr, particulars, details) in enumerate(disc_data, start=1):
        cells = disc_tbl.rows[i].cells
        for cell, text, align in [
            (cells[0], sr,          WD_ALIGN_PARAGRAPH.CENTER),
            (cells[1], particulars, WD_ALIGN_PARAGRAPH.LEFT),
            (cells[2], details,     WD_ALIGN_PARAGRAPH.LEFT),
        ]:
            p = cell.paragraphs[0]
            set_para_fmt(p, align=align, space_before=2, space_after=2)
            add_run(p, text)

    # Page 2 signature block (bottom-right, with signing space above)
    blank_count = 2 if sig_image_path else 4
    for _ in range(blank_count):
        bp = doc.add_paragraph()
        set_para_fmt(bp, space_before=0, space_after=6)
    if sig_image_path:
        sig2_para = doc.add_paragraph()
        set_para_fmt(sig2_para, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
        sig2_para.add_run().add_picture(make_sig_transparent(sig_image_path), width=Cm(2.8))
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
    add_run(p, f'{arb_name}\nARBITRATOR')

    return doc


# ── Document merger ────────────────────────────────────────────────────────────

def merge_noh_docs(doc_bufs):
    """Merge a list of BytesIO DOCX buffers into one Document (page-break separated)."""
    combined = Document(doc_bufs[0])
    for buf in doc_bufs[1:]:
        src  = Document(buf)
        body = combined.element.body
        # Insert a hard page break before appending next document's content
        br_p  = OxmlElement('w:p')
        br_r  = OxmlElement('w:r')
        br_el = OxmlElement('w:br')
        br_el.set(qn('w:type'), 'page')
        br_r.append(br_el)
        br_p.append(br_r)
        last = list(body)[-1]
        if last.tag == qn('w:sectPr'):
            last.addprevious(br_p)
        else:
            body.append(br_p)
        # Copy body elements from source (skip its trailing sectPr)
        for el in list(src.element.body):
            if el.tag == qn('w:sectPr'):
                continue
            el_copy = copy.deepcopy(el)
            last = list(body)[-1]
            if last.tag == qn('w:sectPr'):
                last.addprevious(el_copy)
            else:
                body.append(el_copy)
    return combined


# ── Batch processing ───────────────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    wb = openpyxl.load_workbook(EXCEL_PATH, read_only=True, data_only=True)
    ws = wb[SHEET_NAME]

    all_rows = list(ws.iter_rows(values_only=True))
    headers  = all_rows[0]
    data     = all_rows[1:]
    wb.close()

    total     = 0
    generated = 0
    skipped   = []

    for row_data in data:
        # Skip fully empty rows
        if all(c is None for c in row_data):
            continue
        total += 1

        contract_no   = v(row_data[C_CONTRACT_NO])
        borrower_name = v(row_data[C_BORROWER_NAME])

        # Mandatory field checks
        missing = []
        if not row_data[C_NOH_DATE]:
            missing.append('NOH Date')
        if not v(row_data[C_MEETING_LINK]):
            missing.append('Meeting Link')

        if missing:
            skipped.append({
                'contract_no':   contract_no,
                'borrower_name': borrower_name,
                'reason':        f"Missing: {', '.join(missing)}",
            })
            continue

        try:
            doc = build_noh(row_data)
            safe = contract_no.replace('/', '_').replace('\\', '_')
            filename = f"NOH_{safe}.docx"
            filepath = os.path.join(OUTPUT_DIR, filename)
            doc.save(filepath)
            generated += 1
            print(f"  ✓  {filename}")
        except Exception as e:
            skipped.append({
                'contract_no':   contract_no,
                'borrower_name': borrower_name,
                'reason':        f"Error: {e}",
            })

    # Summary report
    print()
    print("=" * 65)
    print("  BATCH SUMMARY — Notice of Hearing Generator")
    print("=" * 65)
    print(f"  Total rows processed : {total}")
    print(f"  Documents generated  : {generated}")
    print(f"  Rows skipped         : {len(skipped)}")
    if skipped:
        print()
        print(f"  {'Contract No':<25} {'Borrower':<28} Reason")
        print(f"  {'-'*25} {'-'*28} {'-'*28}")
        for s in skipped:
            print(f"  {s['contract_no']:<25} {s['borrower_name']:<28} {s['reason']}")
    print("=" * 65)
    print(f"  Output folder: {OUTPUT_DIR}")
    print()


if __name__ == '__main__':
    main()
